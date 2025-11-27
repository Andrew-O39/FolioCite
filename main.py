from fastapi import FastAPI, Request, Form, Query
from fastapi.responses import HTMLResponse, RedirectResponse, PlainTextResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from db import ensure_default_project, get_projects_for_user, create_project
from io import BytesIO
from docx import Document
from typing import Optional

from starlette.middleware.sessions import SessionMiddleware

from book_citation import (
    format_citation,
    CitationStyle,
    Book,
    format_bibtex,
    format_citation_html,
)
from journal_citation import (
    Article,
    format_article_citation,
    format_article_citation_html,
    format_article_bibtex,
)
from website_citation import (
    Website,
    format_website_citation,
    format_website_citation_html,
    format_website_bibtex,
)
from services import search_books
from journal_services import search_articles
import db
from typing import List, Dict, Any, Optional

from passlib.hash import pbkdf2_sha256
import os
from dotenv import load_dotenv

load_dotenv()  # loads variables from .env into environment

SECRET_KEY = os.getenv("SECRET_KEY")

if not SECRET_KEY:
    # Fallback / safety: you can either raise or use a default for dev
    raise RuntimeError("SECRET_KEY is not set. Please define it in your .env file.")

app = FastAPI(title="FolioCite – Book Citation Generator")
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)


# Mount static files (CSS, logo, etc.)
app.mount("/static", StaticFiles(directory="static"), name="static")

# Jinja2 templates
templates = Jinja2Templates(directory="templates")

db.init_db()

class CurrentUser:
    def __init__(self, user_id: int, username: str, email: str):
        self.id = user_id
        self.username = username
        self.email = email


def get_current_user(request: Request) -> Optional[CurrentUser]:
    user_id = request.session.get("user_id")
    if not user_id:
        return None
    row = db.get_user_by_id(int(user_id))
    if not row:
        return None
    row = db.get_user_by_id(int(user_id))
    return CurrentUser(
        user_id=row["id"],
        username=row["username"],
        email=row["email"],
    )

def get_current_project_id(request: Request, user_id: int) -> int:
    """
    Read the active project_id from the session or create/select the default one.
    """
    project_id = request.session.get("project_id")
    if project_id is not None:
        return int(project_id)

    # Ensure default project exists
    default_project = ensure_default_project(user_id)
    pid = default_project["id"]
    request.session["project_id"] = pid
    return pid

@app.get("/register", response_class=HTMLResponse)
async def register_get(request: Request):
    current_user = get_current_user(request)
    if current_user:
        return RedirectResponse(url="/", status_code=303)

    return templates.TemplateResponse(
        "register.html",
        {"request": request, "error": None, "current_user": current_user},
    )


@app.post("/register", response_class=HTMLResponse)
async def register_post(
    request: Request,
    username: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
    confirm_password: str = Form(...),
):
    # Normalize inputs
    username = username.strip()
    email = email.strip().lower()

    def render_with_error(message: str):
        return templates.TemplateResponse(
            "register.html",
            {
                "request": request,
                "error": message,
                "current_user": None,
                "username": username,
                "email": email,
            },
        )

    if password != confirm_password:
        return render_with_error("Passwords do not match.")

    if db.get_user_by_username(username):
        return render_with_error("Username is already taken.")

    if db.get_user_by_email(email):
        return render_with_error("An account with this email already exists.")

    password_hash = pbkdf2_sha256.hash(password)
    db.create_user(username, email, password_hash)

    # Do NOT log the user in here.
    # Send them to the login page instead.
    return RedirectResponse(url="/login", status_code=303)

@app.get("/login", response_class=HTMLResponse)
async def login_get(request: Request):
    current_user = get_current_user(request)
    if current_user:
        return RedirectResponse(url="/", status_code=303)

    return templates.TemplateResponse(
        "login.html",
        {"request": request, "error": None, "current_user": current_user},
    )


@app.post("/login", response_class=HTMLResponse)
async def login_post(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
):
    username = username.strip()

    def render_with_error(message: str):
        return templates.TemplateResponse(
            "login.html",
            {
                "request": request,
                "error": message,
                "current_user": None,
                "username": username,  # so the form can re-fill it
            },
        )

    # Look up user by username
    row = db.get_user_by_username(username)
    if not row or not pbkdf2_sha256.verify(password, row["password_hash"]):
        return render_with_error("Invalid username or password.")

    # Success: store user_id in session
    request.session["user_id"] = row["id"]
    return RedirectResponse(url="/", status_code=303)


@app.post("/logout", response_class=HTMLResponse)
async def logout(request: Request):
    request.session.clear()
    return RedirectResponse(url="/", status_code=303)

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    return templates.TemplateResponse(
        "index.html",
        {
            "request": request,
            "error": None,
            "selected_style": "apa",
            "selected_source_type": "book",
            "current_user": current_user,
        },
    )

@app.post("/search", response_class=HTMLResponse)
async def search(
    request: Request,
    query: str = Form(...),
    style: str = Form(...),
    source_type: str = Form("book"),  # "book" or "article"
    page: int = Form(1),
):
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    # Clamp page
    if page < 1:
        page = 1

    # ---- Fetch from the appropriate API ----
    if source_type == "article":
        results, had_error = await search_articles(query, limit=10)
    else:
        # default to book
        books, had_error = await search_books(query, limit=10)
        # adapt book results to unified format
        results = []
        for b in books:
            results.append(
                {
                    "entry_type": "book",
                    "title": b.title,
                    "authors": b.authors,
                    "year": b.year,
                    "publisher": b.publisher or "",
                    "place": b.place or "",
                    "journal": "",
                    "volume": "",
                    "issue": "",
                    "pages": "",
                    "doi": "",
                    "cover_url": b.cover_url or "",
                }
            )

    # ---- Error handling ----
    if had_error:
        error_message = (
            "We couldn't reach the external database just now. "
            "This is a temporary problem on their side. Please try again in a few minutes."
        )
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "results": None,
                "error": error_message,
                "selected_style": style,
                "selected_source_type": source_type,
                "current_user": current_user,
            },
        )

    if not results:
        error_message = (
            f"No results found for '{query}'. Please check the spelling or try "
            "a different search (e.g. author name or DOI/ISBN)."
        )
        return templates.TemplateResponse(
            "index.html",
            {
                "request": request,
                "results": None,
                "error": error_message,
                "selected_style": style,
                "selected_source_type": source_type,
                "current_user": current_user,
            },
        )

    # ---- Pagination (5 items per page) ----
    page_size = 5
    total_results = len(results)
    total_pages = (total_results + page_size - 1) // page_size

    # Clamp page again in case the user manually edited it
    if page > total_pages:
        page = total_pages

    start = (page - 1) * page_size
    end = start + page_size
    page_results = results[start:end]

    has_prev = page > 1
    has_next = page < total_pages

    # Render results page
    return templates.TemplateResponse(
        "results.html",
        {
            "request": request,
            "results": page_results,  # only the 5 items for this page
            "style": style,
            "query": query,
            "source_type": source_type,
            "current_page": page,
            "total_pages": total_pages,
            "has_prev": has_prev,
            "has_next": has_next,
            "current_user": current_user,
        },
    )

@app.post("/confirm", response_class=HTMLResponse)
async def confirm(
    request: Request,
    entry_type: str = Form(...),  # "book" or "article"
    title: str = Form(...),
    authors: str = Form(...),
    year: str = Form(None),
    publisher: str = Form(None),
    place: str = Form(None),
    journal: str = Form(None),
    volume: str = Form(None),
    issue: str = Form(None),
    pages: str = Form(None),
    doi: str = Form(None),
    style: str = Form(...),
    cover_url: str = Form(None),
):
    """
    Show a confirmation/edit page where the user can adjust metadata
    before generating a citation.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    return templates.TemplateResponse(
        "confirm.html",
        {
            "request": request,
            "entry_type": entry_type,
            "title": title,
            "authors": authors,
            "year": year or "",
            "publisher": publisher or "",
            "place": place or "",
            "journal": journal or "",
            "volume": volume or "",
            "issue": issue or "",
            "pages": pages or "",
            "doi": doi or "",
            "style": style,
            "cover_url": cover_url or "",
            "current_user": current_user,
        },
    )


@app.get("/manual", response_class=HTMLResponse)
async def manual_entry(
    request: Request,
    source_type: str = "book",   # "book" or "article"
):
    """
    Show a form for manually entering a source (book or journal article).
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    # Ensure only supported types
    if source_type not in ("book", "article"):
        source_type = "book"

    return templates.TemplateResponse(
        "manual.html",
        {
            "request": request,
            "selected_style": "apa",
            "source_type": source_type,
            "current_user": current_user,
        },
    )

@app.post("/cite", response_class=HTMLResponse)
async def cite(
    request: Request,
    entry_type: str = Form(...),   # "book" | "article" | "website"
    title: str = Form(...),
    authors: str = Form(...),
    year: str = Form(None),
    publisher: str = Form(None),
    place: str = Form(None),
    journal: str = Form(None),
    volume: str = Form(None),
    issue: str = Form(None),
    pages: str = Form(None),
    doi: str = Form(None),
    site_name: str = Form(None),
    url: str = Form(None),
    accessed: str = Form(None),
    style: str = Form(...),
    cover_url: str = Form(None),
):
    """
    Generate a formatted citation for the selected source (book, article, or website).
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    authors_list = [a.strip() for a in authors.split(";") if a.strip()]
    chosen_style = CitationStyle(style)

    if entry_type == "article":
        source = Article(
            title=title,
            authors=authors_list,
            year=year or None,
            journal=journal or None,
            volume=volume or None,
            issue=issue or None,
            pages=pages or None,
            doi=doi or None,
        )
        citation = format_article_citation(source, chosen_style)
        citation_html = format_article_citation_html(source, chosen_style)
        bibtex = format_article_bibtex(source)

    elif entry_type == "website":
        source = Website(
            title=title,
            authors=authors_list,
            year=year or None,
            site_name=site_name or None,
            url=url or "",
            accessed=accessed or None,
        )
        citation = format_website_citation(source, chosen_style)
        citation_html = format_website_citation_html(source, chosen_style)
        bibtex = format_website_bibtex(source)

    else:  # default to book
        source = Book(
            title=title,
            authors=authors_list,
            year=year or None,
            publisher=publisher or None,
            place=place or None,
            cover_url=cover_url or None,
        )
        citation = format_citation(source, chosen_style)
        citation_html = format_citation_html(source, chosen_style)
        bibtex = format_bibtex(source)

    return templates.TemplateResponse(
        "citation.html",
        {
            "request": request,
            "entry_type": entry_type,
            "citation": citation,
            "citation_html": citation_html,
            "bibtex": bibtex,
            "style": style,
            "source": source,  # Book, Article or Website
            "current_user": current_user,
        },
    )

@app.get("/bibliography", response_class=HTMLResponse)
async def show_bibliography(
    request: Request,
    style: str = Query("apa"),
    filter_type: str = Query("all"),  # "all" | "book" | "article" | "website"
):
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    # Which project is active?
    project_id = get_current_project_id(request, current_user.id)
    all_projects = get_projects_for_user(current_user.id)

    # Figure out current project name (fallback if something is odd)
    current_project_name = "Current project"
    for p in all_projects:
        if p["id"] == project_id:
            current_project_name = p["name"]
            break

    # Load all rows for this user + project from the DB
    rows = db.get_all_entries(current_user.id, project_id=project_id)

    # Normalise filter_type
    if filter_type not in ("all", "book", "article", "website"):
        filter_type = "all"

    # Apply filter in Python
    if filter_type != "all":
        filtered_rows = [r for r in rows if (r["entry_type"] or "book") == filter_type]
    else:
        filtered_rows = list(rows)

    # Pick citation style (fallback to APA if unknown)
    try:
        chosen_style = CitationStyle(style)
    except ValueError:
        chosen_style = CitationStyle.apa
        style = "apa"

    entries = []

    for row in filtered_rows:
        entry_type = (row["entry_type"] or "book").lower()
        authors_list = [
            a.strip()
            for a in (row["authors"] or "").split(";")
            if a.strip()
        ]

        # Safely pull website-only columns
        row_keys = set(row.keys())
        site_name: Optional[str] = row["site_name"] if "site_name" in row_keys else None
        url: Optional[str] = row["url"] if "url" in row_keys else None
        accessed: Optional[str] = row["accessed"] if "accessed" in row_keys else None

        if entry_type == "article":
            source = Article(
                title=row["title"],
                authors=authors_list,
                year=row["year"],
                journal=row["journal"],
                volume=row["volume"],
                issue=row["issue"],
                pages=row["pages"],
                doi=row["doi"],
            )
            citation = format_article_citation(source, chosen_style)

        elif entry_type == "website":
            source = Website(
                title=row["title"],
                authors=authors_list,
                year=row["year"],
                site_name=site_name,
                url=url or "",
                accessed=accessed,
            )
            citation = format_website_citation(source, chosen_style)

        else:
            # default / legacy: book
            source = Book(
                title=row["title"],
                authors=authors_list,
                year=row["year"],
                publisher=row["publisher"],
                place=row["place"],
                cover_url=row["cover_url"],
            )
            citation = format_citation(source, chosen_style)

        entries.append(
            {
                "id": row["id"],
                "entry_type": entry_type,
                "citation": citation,
                "notes": row["notes"],
            }
        )

    # Sort alphabetically by final citation text
    entries_sorted = sorted(entries, key=lambda e: e["citation"].lower())

    return templates.TemplateResponse(
        "bibliography.html",
        {
            "request": request,
            "entries": entries_sorted,
            "current_style": style,
            "current_filter": filter_type,
            "current_user": current_user,
            "projects": all_projects,
            "current_project_id": project_id,
            "current_project_name": current_project_name,
        },
    )

@app.post("/bibliography/add", response_class=HTMLResponse)
async def add_to_bibliography(
    request: Request,
    entry_type: str = Form(...),
    title: str = Form(...),
    authors: str = Form(...),
    year: str = Form(None),
    publisher: str = Form(None),
    place: str = Form(None),
    journal: str = Form(None),
    volume: str = Form(None),
    issue: str = Form(None),
    pages: str = Form(None),
    doi: str = Form(None),
    site_name: str = Form(None),
    url: str = Form(None),
    accessed: str = Form(None),
    style: str = Form(...),
    cover_url: str = Form(None),
):
    """
    Add a source (book, article, or website) to the persistent bibliography (SQLite).
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    project_id = get_current_project_id(request, current_user.id)

    authors_str = ";".join([a.strip() for a in authors.split(";") if a.strip()])

    db.add_entry(
        current_user.id,
        project_id,
  {
            "entry_type": entry_type,
            "title": title.strip(),
            "authors": authors_str,
            "year": (year or "").strip() or None,
            "publisher": (publisher or "").strip() or None,
            "place": (place or "").strip() or None,
            "journal": (journal or "").strip() or None,
            "volume": (volume or "").strip() or None,
            "issue": (issue or "").strip() or None,
            "pages": (pages or "").strip() or None,
            "doi": (doi or "").strip() or None,
            "site_name": (site_name or "").strip() or None,
            "url": (url or "").strip() or None,
            "accessed": (accessed or "").strip() or None,
            "notes": None,  # initially empty; user can add later
            "style": style,
            "cover_url": (cover_url or "").strip() or None,
        },
    )

    return RedirectResponse(url="/bibliography", status_code=303)


@app.post("/bibliography/clear", response_class=HTMLResponse)
async def clear_bibliography(request: Request):
    """
    Remove all entries from the current user's bibliography.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    project_id = get_current_project_id(request, current_user.id)

    db.clear_entries(current_user.id, project_id=project_id)
    return RedirectResponse(url="/bibliography", status_code=303)

@app.post("/bibliography/notes", response_class=HTMLResponse)
async def update_bibliography_notes(
    request: Request,
    entry_id: int = Form(...),
    notes: str = Form(""),
):
    """
    Update the free-text notes field for a single bibliography entry.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    clean_notes = notes.strip() or None
    db.update_notes(current_user.id, entry_id, clean_notes)

    return RedirectResponse(url="/bibliography", status_code=303)


@app.post("/bibliography/delete", response_class=HTMLResponse)
async def delete_from_bibliography(
    request: Request,
    entry_id: int = Form(...),
):
    """
    Delete a single entry from the current user's bibliography by its ID.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    db.delete_entry(current_user.id, entry_id)
    return RedirectResponse(url="/bibliography", status_code=303)


@app.get("/bibliography/export.txt")
async def export_bibliography_txt(
    request: Request,
    style: str = Query("apa"),
):
    """
    Export all citations as a plain-text .txt file, one per line, in the chosen style.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    rows = db.get_all_entries(current_user.id)

    try:
        chosen_style = CitationStyle(style)
    except ValueError:
        chosen_style = CitationStyle.apa
        style = "apa"

    citations = []
    for row in rows:
        authors_list = [a.strip() for a in (row["authors"] or "").split(";") if a.strip()]
        book = Book(
            title=row["title"],
            authors=authors_list,
            year=row["year"],
            publisher=row["publisher"],
            place=row["place"],
            cover_url=row["cover_url"],
        )
        citations.append(format_citation(book, chosen_style))

    citations_sorted = sorted(citations, key=lambda c: c.lower())
    text = "\n".join(citations_sorted)

    return PlainTextResponse(
        text,
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename=\"bibliography_{style}.txt\"'},
    )


@app.get("/bibliography/export.bib")
async def export_bibliography_bib(
    request: Request,
    style: str = Query("apa"),
):
    """
    Export all entries as a .bib file (BibTeX), in alphabetical order.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    rows = db.get_all_entries(current_user.id)

    bib_entries = []
    for row in rows:
        authors_list = [a.strip() for a in (row["authors"] or "").split(";") if a.strip()]
        book = Book(
            title=row["title"],
            authors=authors_list,
            year=row["year"],
            publisher=row["publisher"],
            place=row["place"],
            cover_url=row["cover_url"],
        )
        bib_entries.append(format_bibtex(book))

    text = "\n\n".join(bib_entries)

    return PlainTextResponse(
        text,
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename=\"bibliography_{style}.bib\"'},
    )

from io import BytesIO
from docx import Document
from fastapi.responses import StreamingResponse

@app.get("/bibliography/export.docx")
async def export_bibliography_docx(
    request: Request,
    style: str = Query("apa"),
    filter_type: str = Query("all"),  # "all" | "book" | "article" | "website"
):
    """
    Export citations (only) as a .docx file that can be opened in Word / Google Docs.
    Notes are intentionally NOT included for consistency with .txt/.bib exports.
    """
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    # Load all rows for this user
    rows = db.get_all_entries(current_user.id)

    # Normalise filter_type
    if filter_type not in ("all", "book", "article", "website"):
        filter_type = "all"

    # Apply filter
    if filter_type != "all":
        filtered_rows = [r for r in rows if (r["entry_type"] or "book") == filter_type]
    else:
        filtered_rows = list(rows)

    # Pick citation style
    try:
        chosen_style = CitationStyle(style)
    except ValueError:
        chosen_style = CitationStyle.apa
        style = "apa"

    # Build list of citations (no notes)
    citations: list[str] = []

    for row in filtered_rows:
        entry_type = (row["entry_type"] or "book").lower()
        authors_list = [
            a.strip()
            for a in (row["authors"] or "").split(";")
            if a.strip()
        ]

        row_keys = set(row.keys())
        site_name = row["site_name"] if "site_name" in row_keys else None
        url = row["url"] if "url" in row_keys else None
        accessed = row["accessed"] if "accessed" in row_keys else None

        if entry_type == "article":
            source = Article(
                title=row["title"],
                authors=authors_list,
                year=row["year"],
                journal=row["journal"],
                volume=row["volume"],
                issue=row["issue"],
                pages=row["pages"],
                doi=row["doi"],
            )
            citation = format_article_citation(source, chosen_style)

        elif entry_type == "website":
            source = Website(
                title=row["title"],
                authors=authors_list,
                year=row["year"],
                site_name=site_name,
                url=url or "",
                accessed=accessed,
            )
            citation = format_website_citation(source, chosen_style)

        else:
            # default: book
            source = Book(
                title=row["title"],
                authors=authors_list,
                year=row["year"],
                publisher=row["publisher"],
                place=row["place"],
                cover_url=row["cover_url"],
            )
            citation = format_citation(source, chosen_style)

        citations.append(citation)

    # Sort alphabetically
    citations_sorted = sorted(citations, key=lambda c: c.lower())

    # Build the .docx document (citations only)
    doc = Document()
    doc.add_heading(f"Bibliography ({style.upper()})", level=1)

    for cit in citations_sorted:
        p = doc.add_paragraph()
        p.add_run(cit)

    # Save to in-memory buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"bibliography_{style}.docx"

    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/projects/select", response_class=HTMLResponse)
async def select_project(
    request: Request,
    project_id: int = Form(...),
):
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    # Make sure this project belongs to the user
    projects = get_projects_for_user(current_user.id)
    if any(p["id"] == project_id for p in projects):
        request.session["project_id"] = int(project_id)

    return RedirectResponse(url="/bibliography", status_code=303)


@app.post("/projects/create", response_class=HTMLResponse)
async def create_project_endpoint(
    request: Request,
    project_name: str = Form(...),
):
    current_user = get_current_user(request)
    if not current_user:
        return RedirectResponse(url="/login", status_code=303)

    name = project_name.strip()
    if not name:
        return RedirectResponse(url="/bibliography", status_code=303)

    # Projects are not "default" – you keep one default per user
    new_project_id = create_project(current_user.id, name, is_default=False)
    request.session["project_id"] = new_project_id

    return RedirectResponse(url="/bibliography", status_code=303)


# For local debugging: `python main.py`
if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
